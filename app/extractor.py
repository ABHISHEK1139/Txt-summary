import fitz
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os
import re
import logging

logger = logging.getLogger(__name__)

# Set tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Set poppler path (relative to app/ folder)
POPPLER_PATH = os.path.join(os.path.dirname(__file__), '..', 'tools', 'poppler-25.12.0', 'Library', 'bin')


def extract_pdf_text(path):
    """Extract text from PDF preserving headings and paragraph structure."""
    try:
        doc = fitz.open(path)
    except Exception as e:
        logger.error(f"Failed to open PDF '{path}': {e}")
        return ""
    
    full_text = []
    
    for page_num, page in enumerate(doc):
        try:
            blocks = page.get_text("dict")["blocks"]
            page_text = []
            
            for block in blocks:
                if block["type"] != 0:  # Skip image blocks
                    continue
                
                block_text = []
                is_heading = False
                
                for line in block["lines"]:
                    line_text = ""
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if not text:
                            continue
                        
                        # Detect headings: bold, large font, or ALL CAPS
                        font_size = span["size"]
                        is_bold = "bold" in span["font"].lower() or "Bold" in span["font"]
                        
                        if font_size > 13 or (is_bold and font_size > 11):
                            is_heading = True
                        
                        line_text += text + " "
                    
                    if line_text.strip():
                        block_text.append(line_text.strip())
                
                if block_text:
                    combined = " ".join(block_text)
                    if is_heading:
                        page_text.append(f"\n\n{combined}\n")
                    else:
                        page_text.append(combined)
            
            if page_text:
                full_text.append("\n\n".join(page_text))
                
        except Exception as e:
            logger.warning(f"Error extracting page {page_num + 1}: {e}")
            # Fallback: try simple text extraction for this page
            try:
                simple_text = page.get_text()
                if simple_text.strip():
                    full_text.append(simple_text)
            except:
                pass
    
    doc.close()
    result = "\n\n".join(full_text)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


def ocr_pdf(path):
    """OCR scanned PDFs using Tesseract + Poppler."""
    # Check if poppler path exists
    if not os.path.isdir(POPPLER_PATH):
        logger.error(f"Poppler not found at: {POPPLER_PATH}")
        logger.error("Install poppler or update POPPLER_PATH in extractor.py")
        return ""
    
    # Check if tesseract is available
    try:
        pytesseract.get_tesseract_version()
    except Exception:
        logger.error("Tesseract not found. Install Tesseract-OCR or update path in extractor.py")
        return ""
    
    try:
        pages = convert_from_path(path, dpi=300, poppler_path=POPPLER_PATH)
        text_parts = []
        for i, page in enumerate(pages):
            try:
                page_text = pytesseract.image_to_string(page)
                if page_text.strip():
                    text_parts.append(page_text.strip())
                logger.info(f"  OCR page {i+1}/{len(pages)} done")
            except Exception as e:
                logger.warning(f"  OCR failed on page {i+1}: {e}")
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"OCR pipeline error: {e}")
        return ""


def extract_docx_text(path):
    """Extract text from Word documents preserving headings and structure."""
    try:
        doc = Document(path)
    except Exception as e:
        logger.error(f"Failed to open DOCX '{path}': {e}")
        return ""
    
    text_parts = []
    
    try:
        for para in doc.paragraphs:
            content = para.text.strip()
            if not content:
                continue
            
            # Detect headings by paragraph style
            style_name = para.style.name.lower() if para.style else ""
            
            if 'heading' in style_name or 'title' in style_name:
                text_parts.append(f"\n\n{content}\n")
            elif 'list' in style_name:
                text_parts.append(f"• {content}")
            else:
                text_parts.append(content)
    except Exception as e:
        logger.warning(f"Error reading paragraphs: {e}")
    
    # Extract text from tables
    try:
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_text.append(" | ".join(row_data))
            if table_text:
                text_parts.append("\n" + "\n".join(table_text) + "\n")
    except Exception as e:
        logger.warning(f"Error reading tables: {e}")
    
    result = "\n\n".join(text_parts)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


def get_text(file_path):
    """Extract text from PDF, DOCX, or TXT files. Fail-safe with clear error messages."""
    
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return ""
    
    file_size = os.path.getsize(file_path)
    if file_size == 0:
        logger.error(f"File is empty: {file_path}")
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Extracting text from '{os.path.basename(file_path)}' ({ext}, {file_size/1024:.1f} KB)")
    
    text = ""
    
    try:
        if ext == ".pdf":
            text = extract_pdf_text(file_path)
            # If empty or very short, assume scanned PDF and try OCR
            if len(text.strip()) < 200:
                logger.info("PDF text extraction yielded minimal results, trying OCR...")
                ocr_text = ocr_pdf(file_path)
                if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    logger.info(f"OCR extracted {len(text)} characters")
                elif not text.strip():
                    logger.warning("Both PDF extraction and OCR failed to get text")
            else:
                logger.info(f"PDF extracted {len(text)} characters")
                
        elif ext == ".docx":
            text = extract_docx_text(file_path)
            logger.info(f"DOCX extracted {len(text)} characters")
            
        elif ext in (".txt", ".text", ".md"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            logger.info(f"Text file read {len(text)} characters")
            
        elif ext == ".doc":
            logger.warning(".doc format not fully supported. Please convert to .docx")
            return ""
        else:
            logger.warning(f"Unsupported file type: {ext}")
            return ""
            
    except Exception as e:
        logger.error(f"Unexpected error extracting text: {e}")
        return ""
    
    # Final cleanup
    if text:
        text = text.strip()
        word_count = len(text.split())
        logger.info(f"✅ Extraction complete: {word_count} words, {len(text)} chars")
    else:
        logger.warning("⚠️ No text could be extracted from the file")
    
    return text
