"""End-to-end test of the full pipeline: extraction, chunking, OCR, poppler."""
import os

passed = 0
failed = 0

def test(name, condition, detail=""):
    global passed, failed
    if condition:
        passed += 1
        print(f"  ✅ {name}")
    else:
        failed += 1
        print(f"  ❌ {name}: {detail}")

# TEST 1: PDF Extraction (PyMuPDF)
print("=" * 50)
print("TEST 1: PyMuPDF (PDF text extraction)")
print("=" * 50)
try:
    import fitz
    test("PyMuPDF imported", True)
    test("Version", True, fitz.__doc__[:30])
except ImportError as e:
    test("PyMuPDF import", False, str(e))

# TEST 2: Tesseract OCR
print("\n" + "=" * 50)
print("TEST 2: Tesseract OCR")
print("=" * 50)
try:
    import pytesseract
    v = pytesseract.get_tesseract_version()
    test("Tesseract found", True)
    print(f"       Version: {v}")
except Exception as e:
    test("Tesseract", False, str(e))

# TEST 3: Poppler
print("\n" + "=" * 50)
print("TEST 3: Poppler (PDF to image)")
print("=" * 50)
from src.extractor import POPPLER_PATH
test("Poppler dir exists", os.path.isdir(POPPLER_PATH))
pdftoppm = os.path.join(POPPLER_PATH, 'pdftoppm.exe')
test("pdftoppm.exe exists", os.path.isfile(pdftoppm))
print(f"       Path: {POPPLER_PATH}")

# TEST 4: DOCX Extraction
print("\n" + "=" * 50)
print("TEST 4: DOCX Extraction")
print("=" * 50)
try:
    from docx import Document
    doc = Document()
    doc.add_heading('Test Heading', level=1)
    doc.add_paragraph('This is a test paragraph with enough content to verify extraction works correctly.')
    doc.add_heading('Second Section', level=2)
    doc.add_paragraph('This paragraph is under the second heading for multi-section testing.')
    test_path = 'uploads/test_doc.docx'
    os.makedirs('uploads', exist_ok=True)
    doc.save(test_path)
    
    from src.extractor import get_text
    text = get_text(test_path)
    test("DOCX extracted text", len(text) > 50)
    has_heading = 'Test Heading' in text
    test("Headings preserved", has_heading)
    has_content = 'test paragraph' in text
    test("Content preserved", has_content)
    print(f"       Words: {len(text.split())}")
    print(f"       Preview: {text[:100]}...")
    os.remove(test_path)
except Exception as e:
    test("DOCX extraction", False, str(e))

# TEST 5: TXT Extraction
print("\n" + "=" * 50)
print("TEST 5: TXT Extraction")
print("=" * 50)
try:
    test_txt = 'uploads/test.txt'
    with open(test_txt, 'w') as f:
        f.write('This is a simple text file test.\nIt has multiple lines.\nThird line here.')
    text = get_text(test_txt)
    test("TXT read", len(text) > 20)
    test("Content correct", 'simple text file' in text)
    os.remove(test_txt)
except Exception as e:
    test("TXT extraction", False, str(e))

# TEST 6: Chunking & Sections
print("\n" + "=" * 50)
print("TEST 6: Context-Aware Chunking")
print("=" * 50)
try:
    from src.summarizer import _split_into_sections, _extract_doc_overview, _is_heading
    
    sample = """Introduction
This is the introduction paragraph that explains the background of the research. It covers multiple topics and provides context.

Methodology
The methodology section describes how the experiments were conducted. Various tools and techniques were used.

Results
The results show significant improvement over the baseline approach. Key findings are presented in this section."""

    chunks = _split_into_sections(sample, max_words=30)
    test("Sections detected", len(chunks) >= 2)
    print(f"       Sections found: {len(chunks)}")
    
    overview = _extract_doc_overview(sample)
    test("Overview extracted", len(overview) > 10)
    print(f"       Overview: {overview[:80]}")
    
    test("Heading detection: 'Introduction'", _is_heading('Introduction'))
    test("Heading detection: 'RESULTS'", _is_heading('RESULTS'))
    test("Not heading: long sentence", not _is_heading('This is a regular sentence that ends with a period.'))
except Exception as e:
    test("Chunking", False, str(e))

# SUMMARY
print("\n" + "=" * 50)
print(f"RESULTS: {passed} passed, {failed} failed")
print("=" * 50)
if failed == 0:
    print("🎉 ALL TESTS PASSED!")
else:
    print(f"⚠️ {failed} test(s) need attention")
